import cv2
import torch
import lpips
from PIL import Image
import torchvision.transforms as transforms
import time
import os
import easyocr
from docx import Document
from datetime import datetime
from docx.shared import Inches
import csv

LPIPS_THRESHOLD = 0.045
CHECK_INTERVAL = 2
SAVE_FOLDER = "detected_changes"

CSV_FILE = os.path.join(SAVE_FOLDER, "zone1_detections.csv")

# Initialize EasyOCR reader once at startup (expensive to create repeatedly)
# gpu=True if you have a CUDA GPU, otherwise False
print("Loading EasyOCR model...")
reader = easyocr.Reader(['en'], gpu=torch.cuda.is_available())
print("EasyOCR ready.")

if not os.path.exists(SAVE_FOLDER):
    os.makedirs(SAVE_FOLDER)


# ── CSV helpers ──────────────────────────────────────────────────────────────
def ensure_csv_header(csv_path: str):
    if not os.path.exists(csv_path) or os.path.getsize(csv_path) == 0:
        with open(csv_path, mode="w", newline="", encoding="utf-8") as f:
            writer = csv.writer(f)
            writer.writerow(["Zone 1", "Timestamp", "LPIPS_Distance", "Extracted_Text", "ROI_Image_Path"])


def append_to_csv(csv_path: str, timestamp: str, lpips_dist: float, extracted_text: str, roi_image_path: str):
    with open(csv_path, mode="a", newline="", encoding="utf-8") as f:
        writer = csv.writer(f)
        writer.writerow(["Zone 1", timestamp, f"{lpips_dist:.6f}", extracted_text.strip(), roi_image_path])


ensure_csv_header(CSV_FILE)


# ── LPIPS model ───────────────────────────────────────────────────────────────
loss_fn = lpips.LPIPS(net='alex')
loss_fn.eval()
device = torch.device('cuda' if torch.cuda.is_available() else 'cpu')
loss_fn = loss_fn.to(device)


# ── ROI state ─────────────────────────────────────────────────────────────────
x1 = y1 = 0
x2 = y2 = 0

drag_left = drag_right = drag_top = drag_bottom = False
roi_moved = False

BTN_X1, BTN_Y1 = 10, 10
BTN_X2, BTN_Y2 = 140, 50

LINE_GRAB_PX = 10
MIN_ROI_SIZE = 30


def clamp(v, lo, hi):
    return max(lo, min(hi, v))


def normalize_roi(x1_, y1_, x2_, y2_, w, h):
    x1_ = clamp(x1_, 0, w - 1)
    x2_ = clamp(x2_, 0, w - 1)
    y1_ = clamp(y1_, 0, h - 1)
    y2_ = clamp(y2_, 0, h - 1)

    if x1_ > x2_:
        x1_, x2_ = x2_, x1_
    if y1_ > y2_:
        y1_, y2_ = y2_, y1_

    if (x2_ - x1_) < MIN_ROI_SIZE:
        x2_ = clamp(x1_ + MIN_ROI_SIZE, 0, w - 1)
        x1_ = clamp(x2_ - MIN_ROI_SIZE, 0, w - 1)

    if (y2_ - y1_) < MIN_ROI_SIZE:
        y2_ = clamp(y1_ + MIN_ROI_SIZE, 0, h - 1)
        y1_ = clamp(y2_ - MIN_ROI_SIZE, 0, h - 1)

    return x1_, y1_, x2_, y2_


def cv2_to_tensor(img_bgr):
    img = cv2.cvtColor(img_bgr, cv2.COLOR_BGR2RGB)
    pil_img = Image.fromarray(img)
    transform = transforms.Compose([
        transforms.Resize((256, 256)),
        transforms.ToTensor(),
        transforms.Normalize((0.5, 0.5, 0.5), (0.5, 0.5, 0.5))
    ])
    return transform(pil_img).unsqueeze(0).to(device)


# ── OCR with EasyOCR ──────────────────────────────────────────────────────────
def preprocess_for_ocr(image_bgr):
    """Upscale + threshold to help EasyOCR with glowing/LED-style digits."""
    # Upscale 4x — dramatically improves accuracy on small ROIs
    upscaled = cv2.resize(image_bgr, None, fx=4, fy=4, interpolation=cv2.INTER_CUBIC)

    # Convert to grayscale and boost contrast
    gray = cv2.cvtColor(upscaled, cv2.COLOR_BGR2GRAY)
    clahe = cv2.createCLAHE(clipLimit=3.0, tileGridSize=(8, 8))
    gray = clahe.apply(gray)

    # Threshold: isolate bright digits on dark background
    _, thresh = cv2.threshold(gray, 160, 255, cv2.THRESH_BINARY)

    # Close small gaps caused by glow/bloom
    kernel = cv2.getStructuringElement(cv2.MORPH_RECT, (2, 2))
    thresh = cv2.morphologyEx(thresh, cv2.MORPH_CLOSE, kernel)

    # EasyOCR expects BGR or RGB; convert back to BGR 3-channel
    return cv2.cvtColor(thresh, cv2.COLOR_GRAY2BGR)


def extract_text_from_image(image_path):
    """Extract text from saved ROI image using EasyOCR."""
    image = cv2.imread(image_path)
    preprocessed = preprocess_for_ocr(image)

    # allowlist keeps it focused on numbers + decimal point
    results = reader.readtext(
        preprocessed,
        allowlist='0123456789.',
        detail=0,          # return text only, not bounding boxes
        paragraph=False    # treat each detection independently
    )

    text = ' '.join(results).strip()
    return text if text else "[no text detected]"


# ── Word export ───────────────────────────────────────────────────────────────
def write_to_word(timestamp, text, roi_image_path):
    doc = Document()
    doc.add_heading("Camera Change Detection (ROI)", level=1)
    doc.add_paragraph(f"Time detected: {timestamp}")
    doc.add_paragraph("Detected ROI Image:")
    doc.add_picture(roi_image_path, width=Inches(5.8))
    doc.add_paragraph("\nExtracted Text (ROI):")
    doc.add_paragraph(text)

    filename = os.path.join(SAVE_FOLDER, f"Camera_Detection_{timestamp.replace(':', '-')}.docx")
    doc.save(filename)
    print(f">>> Word document saved as: {filename}")


# ── ROI helpers ───────────────────────────────────────────────────────────────
def reset_roi_to_center(w, h):
    global x1, y1, x2, y2, roi_moved
    cx, cy = w // 2, h // 2
    x1, y1 = cx - w // 6, cy - h // 6
    x2, y2 = cx + w // 6, cy + h // 6
    x1, y1, x2, y2 = normalize_roi(x1, y1, x2, y2, w, h)
    roi_moved = True


def on_mouse(event, mx, my, flags, param):
    global x1, y1, x2, y2
    global drag_left, drag_right, drag_top, drag_bottom
    global roi_moved

    w, h = param["w"], param["h"]

    if event == cv2.EVENT_LBUTTONDOWN:
        if BTN_X1 <= mx <= BTN_X2 and BTN_Y1 <= my <= BTN_Y2:
            reset_roi_to_center(w, h)
            drag_left = drag_right = drag_top = drag_bottom = False
            print("ROI reset.")
            return

        d_left   = abs(mx - x1)
        d_right  = abs(mx - x2)
        d_top    = abs(my - y1)
        d_bottom = abs(my - y2)

        candidates = [("left", d_left), ("right", d_right), ("top", d_top), ("bottom", d_bottom)]
        candidates.sort(key=lambda t: t[1])
        closest_name, closest_dist = candidates[0]

        if closest_dist <= LINE_GRAB_PX:
            drag_left   = closest_name == "left"
            drag_right  = closest_name == "right"
            drag_top    = closest_name == "top"
            drag_bottom = closest_name == "bottom"

    elif event == cv2.EVENT_MOUSEMOVE:
        changed = False
        if drag_left:   x1 = mx; changed = True
        if drag_right:  x2 = mx; changed = True
        if drag_top:    y1 = my; changed = True
        if drag_bottom: y2 = my; changed = True

        if changed:
            x1, y1, x2, y2 = normalize_roi(x1, y1, x2, y2, w, h)
            roi_moved = True

    elif event == cv2.EVENT_LBUTTONUP:
        drag_left = drag_right = drag_top = drag_bottom = False


# ── Main loop ─────────────────────────────────────────────────────────────────
cap = cv2.VideoCapture(0)
if not cap.isOpened():
    raise RuntimeError("Cannot open camera")

WINDOW_NAME = "Camera Feed"
cv2.namedWindow(WINDOW_NAME)

ret, frame = cap.read()
if not ret:
    raise RuntimeError("Failed to grab initial frame")

h, w = frame.shape[:2]
reset_roi_to_center(w, h)
cv2.setMouseCallback(WINDOW_NAME, on_mouse, param={"w": w, "h": h})

last_check = 0
prev_tensor = None

print("System running... press 'q' to quit. Press 'r' to reset ROI.")
print("Drag 4 lines: RED = left/right, YELLOW = top/bottom.")
print(f"CSV logging to: {CSV_FILE}")

while True:
    ret, frame = cap.read()
    if not ret:
        print("Failed to grab frame")
        break

    h, w = frame.shape[:2]
    x1, y1, x2, y2 = normalize_roi(x1, y1, x2, y2, w, h)

    overlay = frame.copy()

    # Reset button
    cv2.rectangle(overlay, (BTN_X1, BTN_Y1), (BTN_X2, BTN_Y2), (40, 40, 40), -1)
    cv2.rectangle(overlay, (BTN_X1, BTN_Y1), (BTN_X2, BTN_Y2), (255, 255, 255), 2)
    cv2.putText(overlay, "RESET", (BTN_X1 + 20, BTN_Y1 + 28),
                cv2.FONT_HERSHEY_SIMPLEX, 0.7, (255, 255, 255), 2)

    # ROI lines
    cv2.line(overlay, (0, y1), (w - 1, y1), (0, 255, 255), 2)
    cv2.line(overlay, (0, y2), (w - 1, y2), (0, 255, 255), 2)
    cv2.line(overlay, (x1, 0), (x1, h - 1), (0, 0, 255), 2)
    cv2.line(overlay, (x2, 0), (x2, h - 1), (0, 0, 255), 2)
    cv2.rectangle(overlay, (x1, y1), (x2, y2), (255, 255, 255), 2)
    cv2.putText(overlay, f"ROI: ({x1},{y1})-({x2},{y2})  {x2-x1}x{y2-y1}",
                (10, 80), cv2.FONT_HERSHEY_SIMPLEX, 0.6, (255, 255, 255), 2)

    cv2.imshow(WINDOW_NAME, overlay)

    roi = frame[y1:y2, x1:x2]
    if roi.size == 0 or roi.shape[0] < MIN_ROI_SIZE or roi.shape[1] < MIN_ROI_SIZE:
        key = cv2.waitKey(1) & 0xFF
        if key == ord('q'):
            break
        if key == ord('r'):
            reset_roi_to_center(w, h)
            prev_tensor = None
        continue

    if roi_moved:
        prev_tensor = None
        roi_moved = False
        print("ROI moved -> baseline reset.")

    if time.time() - last_check >= CHECK_INTERVAL:
        last_check = time.time()
        curr_tensor = cv2_to_tensor(roi)

        if prev_tensor is not None:
            dist = loss_fn(prev_tensor, curr_tensor).item()
            print(f"LPIPS Distance (ROI): {dist:.3f}")

            if dist > LPIPS_THRESHOLD:
                print(">>> CHANGE DETECTED (ROI) <<<")

                timestamp = datetime.now().strftime("%Y-%m-%d %H-%M-%S")
                roi_image_file = os.path.join(SAVE_FOLDER, f"roi_change_{timestamp}.jpg")

                cv2.imwrite(roi_image_file, roi)
                print(f"Saved ROI frame: {roi_image_file}")

                extracted_text = extract_text_from_image(roi_image_file)
                print(f"Extracted text (ROI): {extracted_text}")

                append_to_csv(CSV_FILE, timestamp, dist, extracted_text, roi_image_file)
                write_to_word(timestamp, extracted_text, roi_image_file)
            else:
                print("No significant change (ROI)")

        prev_tensor = curr_tensor

    key = cv2.waitKey(1) & 0xFF
    if key == ord('q'):
        break
    if key == ord('r'):
        reset_roi_to_center(w, h)
        prev_tensor = None
        print("ROI reset.")

cap.release()
cv2.destroyAllWindows()