import cv2
import torch
import lpips
from PIL import Image
import torchvision.transforms as transforms
import time
import os
import pytesseract
from docx import Document
from datetime import datetime
import requests

# ----------- CONFIGURATION----------- #
LPIPS_THRESHOLD = 0.1
CHECK_INTERVAL = 5
SAVE_FOLDER = "detected_changes"
os.makedirs(SAVE_FOLDER, exist_ok=True)

pytesseract.pytesseract.tesseract_cmd = r"C:\Users\G00420041@atu.ie\AppData\Local\Programs\Tesseract-OCR\tesseract.exe"

# ----------- LPIPS MODEL ----------- #
loss_fn = lpips.LPIPS(net='alex')
loss_fn.eval()
device = torch.device('cuda' if torch.cuda.is_available() else 'cpu')
loss_fn = loss_fn.to(device)

# ----------- UTILITY FUNCTIONS ----------- #
def cv2_to_tensor(img):
    img = cv2.cvtColor(img, cv2.COLOR_BGR2RGB)
    pil_img = Image.fromarray(img)
    transform = transforms.Compose([
        transforms.Resize((256, 256)),
        transforms.ToTensor(),
        transforms.Normalize((0.5, 0.5, 0.5),(0.5, 0.5, 0.5))
    ])
    return transform(pil_img).unsqueeze(0).to(device)

def extract_text_from_image(image_path):
    image = cv2.imread(image_path)
    gray = cv2.cvtColor(image, cv2.COLOR_BGR2GRAY)
    gray = cv2.threshold(gray, 0,255,cv2.THRESH_BINARY | cv2.THRESH_OTSU)[1]
    temp_file = "temp_ocr.jpg"
    cv2.imwrite(temp_file, gray)
    text = pytesseract.image_to_string(Image.open(temp_file))
    os.remove(temp_file)
    return text

def write_to_word(timestamp, text):
    filename = f"Camera_Detection_{timestamp.replace(':','-')}.docx"
    doc = Document()
    doc.add_heading("Camera Change Detection", level=1)
    doc.add_paragraph(f"Time detected: {timestamp}")
    doc.add_paragraph("Extracted text:")
    doc.add_paragraph(text)
    doc.save(filename)
    print(f">>> Word document saved as: {filename}")
    return filename

def upload_word_doc(file_path):
    url = "http://127.0.0.1:8000/upload_word/"
    with open(file_path, "rb") as f:
        files = {"file": (os.path.basename(file_path), f, "application/vnd.openxmlformats-officedocument.wordprocessingml.document")}
        response = requests.post(url, files=files)
    print("Upload response:", response.json())

# ----------- MY CAMERA SETUP ----------- #
cap = cv2.VideoCapture(0)
if not cap.isOpened():
    raise RuntimeError("Cannot open camera")

last_check = 0
prev_tensor = None

print("System running... press 'q' to quit.")

# ----------- MAIN LOOP ----------- #
while True:
    ret, frame = cap.read()
    if not ret:
        print("Failed to grab frame")
        break

    cv2.imshow("Camera Feed", frame)

    if time.time() - last_check >= CHECK_INTERVAL:
        last_check = time.time()
        curr_tensor = cv2_to_tensor(frame)

        if prev_tensor is not None:
            dist = loss_fn(prev_tensor, curr_tensor).item()
            print(f"LPIPS Distance: {dist:.3f}")

            if dist > LPIPS_THRESHOLD:
                print(">>> CHANGE DETECTED <<<")
                timestamp = datetime.now().strftime("%Y-%m-%d %H-%M-%S")
                image_file = os.path.join(SAVE_FOLDER, f"change_{timestamp}.jpg")
                cv2.imwrite(image_file, frame)
                print(f"Saved changed frame: {image_file}")

                extracted_text = extract_text_from_image(image_file)
                print("Extracted text:\n", extracted_text)

                word_file = write_to_word(timestamp, extracted_text)
                upload_word_doc(word_file)
            else:
                print("No significant change")

        prev_tensor = curr_tensor

    if cv2.waitKey(1) & 0xFF == ord('q'):
        break

cap.release()
cv2.destroyAllWindows()
