import cv2
import torch
import lpips
from PIL import Image
import torchvision.transforms as transforms
import time
import os
import pytesseract
from docx import Document
from datetime import datetime
from docx.shared import Inches
import csv
import numpy as np

LPIPS_THRESHOLD = 0.045
CHECK_INTERVAL = 2
SAVE_FOLDER = "detected_changes"
CSV_FILE = os.path.join(SAVE_FOLDER, "zone1_detections.csv")

# Turn ON to confirm files are saving even when LPIPS doesn't trigger
DEBUG_SAVE_PREPROC_EVERY_CHECK = False

pytesseract.pytesseract.tesseract_cmd = r"C:\Program Files\Tesseract-OCR\tesseract.exe"

if not os.path.exists(SAVE_FOLDER):
    os.makedirs(SAVE_FOLDER)

# CSV SETUP 
def ensure_csv_header(csv_path: str):
    if not os.path.exists(csv_path) or os.path.getsize(csv_path) == 0:
        with open(csv_path, mode="w", newline="", encoding="utf-8") as f:
            writer = csv.writer(f)
            writer.writerow([
                "Zone 1", "Timestamp", "LPIPS_Distance", "Extracted_Text",
                "ROI_Raw_Path", "ROI_Preproc_Path"
            ])

def append_to_csv(csv_path: str, timestamp: str, lpips_dist: float,
                  extracted_text: str, roi_raw_path: str, roi_preproc_path: str):
    with open(csv_path, mode="a", newline="", encoding="utf-8") as f:
        writer = csv.writer(f)
        writer.writerow([
            "Zone 1",
            timestamp,
            f"{lpips_dist:.6f}",
            extracted_text.strip(),
            roi_raw_path,
            roi_preproc_path
        ])

ensure_csv_header(CSV_FILE)







# LPIPS MODEL 
loss_fn = lpips.LPIPS(net='alex')
loss_fn.eval()
device = torch.device('cuda' if torch.cuda.is_available() else 'cpu')
loss_fn = loss_fn.to(device)






# ROI STATE 
x1 = y1 = 0
x2 = y2 = 0

drag_left = drag_right = drag_top = drag_bottom = False
roi_moved = False

# Reset button
BTN_X1, BTN_Y1 = 10, 10
BTN_X2, BTN_Y2 = 140, 50
LINE_GRAB_PX = 10
MIN_ROI_SIZE = 30

# HELPERS 
def clamp(v, lo, hi):
    return max(lo, min(hi, v))

def normalize_roi(x1_, y1_, x2_, y2_, w, h):
    """Ensure x1<x2 and y1<y2, keep inside frame, enforce minimum size."""
    x1_ = clamp(x1_, 0, w - 1)
    x2_ = clamp(x2_, 0, w - 1)
    y1_ = clamp(y1_, 0, h - 1)
    y2_ = clamp(y2_, 0, h - 1)

    if x1_ > x2_:
        x1_, x2_ = x2_, x1_
    if y1_ > y2_:
        y1_, y2_ = y2_, y1_

    if (x2_ - x1_) < MIN_ROI_SIZE:
        x2_ = clamp(x1_ + MIN_ROI_SIZE, 0, w - 1)
        x1_ = clamp(x2_ - MIN_ROI_SIZE, 0, w - 1)

    if (y2_ - y1_) < MIN_ROI_SIZE:
        y2_ = clamp(y1_ + MIN_ROI_SIZE, 0, h - 1)
        y1_ = clamp(y2_ - MIN_ROI_SIZE, 0, h - 1)

    return x1_, y1_, x2_, y2_

def cv2_to_tensor(img_bgr):
    """Convert BGR cv2 image -> LPIPS tensor."""
    img = cv2.cvtColor(img_bgr, cv2.COLOR_BGR2RGB)
    pil_img = Image.fromarray(img)
    transform = transforms.Compose([
        transforms.Resize((256, 256)),
        transforms.ToTensor(),
        transforms.Normalize((0.5, 0.5, 0.5), (0.5, 0.5, 0.5))
    ])
    return transform(pil_img).unsqueeze(0).to(device)

# OCR PREPROCESSING (tuned for bright digits on blue/green UI) 
def preprocess_roi_for_ocr(roi_bgr: np.ndarray) -> np.ndarray:
    """
    For UI overlays like your sample (bright digits on teal/blue with grid lines):
    - Upscale
    - Use LAB L-channel
    - TOPHAT to highlight bright text
    - Otsu binarize
    - Remove horizontal grid lines
    - Morph close
    """
    if roi_bgr is None or roi_bgr.size == 0:
        return None

    # Upscale helps a lot for small digits
    roi = cv2.resize(roi_bgr, None, fx=4.0, fy=4.0, interpolation=cv2.INTER_CUBIC)

    # L channel = brightness, robust vs color backgrounds
    lab = cv2.cvtColor(roi, cv2.COLOR_BGR2LAB)
    L = lab[:, :, 0]

    # Contrast boost
    clahe = cv2.createCLAHE(clipLimit=3.0, tileGridSize=(8, 8))
    L = clahe.apply(L)

    # TOPHAT emphasizes bright text over darker background
    tophat_kernel = cv2.getStructuringElement(cv2.MORPH_RECT, (25, 25))
    tophat = cv2.morphologyEx(L, cv2.MORPH_TOPHAT, tophat_kernel)

    tophat = cv2.GaussianBlur(tophat, (5, 5), 0)

    bw = cv2.threshold(tophat, 0, 255, cv2.THRESH_BINARY + cv2.THRESH_OTSU)[1]

    # Remove horizontal grid lines (common on camera overlays)
    hkernel = cv2.getStructuringElement(cv2.MORPH_RECT, (80, 1))
    lines = cv2.morphologyEx(bw, cv2.MORPH_OPEN, hkernel, iterations=2)
    bw = cv2.subtract(bw, lines)

    # Clean up
    bw = cv2.morphologyEx(
        bw, cv2.MORPH_CLOSE,
        cv2.getStructuringElement(cv2.MORPH_RECT, (3, 3)),
        iterations=1
    )

    return bw

def crop_to_text_components(bw: np.ndarray) -> np.ndarray:
    """
    Crop the binary image to the biggest connected components so OCR ignores noise.
    """
    if bw is None:
        return None

    contours, _ = cv2.findContours(bw, cv2.RETR_EXTERNAL, cv2.CHAIN_APPROX_SIMPLE)
    if not contours:
        return bw

    areas = [cv2.contourArea(c) for c in contours]
    if not areas:
        return bw

    # Keep only meaningful blobs
    big = [c for c in contours if cv2.contourArea(c) > 500]
    if not big:
        return bw

    h, w = bw.shape[:2]
    xs, ys, x2s, y2s = [], [], [], []
    for c in big:
        x, y, ww, hh = cv2.boundingRect(c)
        xs.append(x); ys.append(y); x2s.append(x + ww); y2s.append(y + hh)

    pad = 10
    x1 = max(min(xs) - pad, 0)
    y1 = max(min(ys) - pad, 0)
    x2 = min(max(x2s) + pad, w)
    y2 = min(max(y2s) + pad, h)

    return bw[y1:y2, x1:x2]

def ocr_from_preprocessed(preproc_bw: np.ndarray) -> str:
    """
    Tesseract config focused on a single line of digits with optional decimal point.
    """
    if preproc_bw is None:
        return ""

    # Crop to main text blobs
    preproc_bw = crop_to_text_components(preproc_bw)

    # whitelist includes '.' for values like 4.5
    config = "--oem 3 --psm 7 -c tessedit_char_whitelist=0123456789."
    pil_img = Image.fromarray(preproc_bw)
    text = pytesseract.image_to_string(pil_img, config=config)
    return text

def extract_text_and_save_preproc(roi_bgr: np.ndarray, save_base_path_no_ext: str):
    """
    Saves BOTH:
      - raw ROI:         <base>_raw.jpg
      - preprocessed:    <base>_preproc.png
    Returns: (text, raw_path, preproc_path)
    """
    raw_path = save_base_path_no_ext + "_raw.jpg"
    preproc_path = save_base_path_no_ext + "_preproc.png"

    cv2.imwrite(raw_path, roi_bgr)

    preproc = preprocess_roi_for_ocr(roi_bgr)
    if preproc is not None:
        # also save the cropped-to-text version (optional but useful)
        cropped = crop_to_text_components(preproc)
        cv2.imwrite(preproc_path, cropped if cropped is not None else preproc)
        text = ocr_from_preprocessed(preproc)
    else:
        text = ""
        preproc_path = ""

    return text, raw_path, preproc_path

# REPORTING 
def write_to_word(timestamp, text, roi_image_path):
    doc = Document()
    doc.add_heading("Camera Change Detection (ROI)", level=1)

    doc.add_paragraph(f"Time detected: {timestamp}")

    doc.add_paragraph("Detected ROI Image:")
    doc.add_picture(roi_image_path, width=Inches(5.8))

    doc.add_paragraph("\nExtracted Text (ROI):")
    doc.add_paragraph(text)

    filename = os.path.join(SAVE_FOLDER, f"Camera_Detection_{timestamp.replace(':','-')}.docx")
    doc.save(filename)
    print(f">>> Word document saved as: {filename}")

def reset_roi_to_center(w, h):
    """Helper: center ROI and mark moved."""
    global x1, y1, x2, y2, roi_moved
    cx, cy = w // 2, h // 2
    x1, y1 = cx - w // 6, cy - h // 6
    x2, y2 = cx + w // 6, cy + h // 6
    x1, y1, x2, y2 = normalize_roi(x1, y1, x2, y2, w, h)
    roi_moved = True

def on_mouse(event, mx, my, flags, param):
    global x1, y1, x2, y2
    global drag_left, drag_right, drag_top, drag_bottom
    global roi_moved

    w, h = param["w"], param["h"]

    if event == cv2.EVENT_LBUTTONDOWN:
        if BTN_X1 <= mx <= BTN_X2 and BTN_Y1 <= my <= BTN_Y2:
            reset_roi_to_center(w, h)
            drag_left = drag_right = drag_top = drag_bottom = False
            print("ROI reset.")
            return

        d_left = abs(mx - x1)
        d_right = abs(mx - x2)
        d_top = abs(my - y1)
        d_bottom = abs(my - y2)

        candidates = [("left", d_left), ("right", d_right), ("top", d_top), ("bottom", d_bottom)]
        candidates.sort(key=lambda t: t[1])
        closest_name, closest_dist = candidates[0]

        if closest_dist <= LINE_GRAB_PX:
            drag_left = closest_name == "left"
            drag_right = closest_name == "right"
            drag_top = closest_name == "top"
            drag_bottom = closest_name == "bottom"

    elif event == cv2.EVENT_MOUSEMOVE:
        changed = False

        if drag_left:
            x1 = mx
            changed = True
        if drag_right:
            x2 = mx
            changed = True
        if drag_top:
            y1 = my
            changed = True
        if drag_bottom:
            y2 = my
            changed = True

        if changed:
            x1_, y1_, x2_, y2_ = normalize_roi(x1, y1, x2, y2, w, h)
            x1, y1, x2, y2 = x1_, y1_, x2_, y2_
            roi_moved = True

    elif event == cv2.EVENT_LBUTTONUP:
        drag_left = drag_right = drag_top = drag_bottom = False



# CAMERA SETUP 
cap = cv2.VideoCapture(0)
if not cap.isOpened():
    raise RuntimeError("Cannot open camera")

WINDOW_NAME = "Camera Feed"
cv2.namedWindow(WINDOW_NAME)

ret, frame = cap.read()
if not ret:
    raise RuntimeError("Failed to grab initial frame")

h, w = frame.shape[:2]
reset_roi_to_center(w, h)
cv2.setMouseCallback(WINDOW_NAME, on_mouse, param={"w": w, "h": h})

last_check = 0
prev_tensor = None

print("System running... press 'q' to quit. Press 'r' to reset ROI.")
print("Drag 4 lines: RED = left/right, YELLOW = top/bottom.")
print("CSV logging to:", CSV_FILE)





# MAIN LOOP 
while True:
    ret, frame = cap.read()
    if not ret:
        print("Failed to grab frame")
        break

    h, w = frame.shape[:2]
    x1, y1, x2, y2 = normalize_roi(x1, y1, x2, y2, w, h)

    overlay = frame.copy()

    # reset button
    cv2.rectangle(overlay, (BTN_X1, BTN_Y1), (BTN_X2, BTN_Y2), (40, 40, 40), -1)
    cv2.rectangle(overlay, (BTN_X1, BTN_Y1), (BTN_X2, BTN_Y2), (255, 255, 255), 2)
    cv2.putText(overlay, "RESET", (BTN_X1 + 20, BTN_Y1 + 28),
                cv2.FONT_HERSHEY_SIMPLEX, 0.7, (255, 255, 255), 2)

    # top and bottom borders in yellow
    cv2.line(overlay, (0, y1), (w - 1, y1), (0, 255, 255), 2)
    cv2.line(overlay, (0, y2), (w - 1, y2), (0, 255, 255), 2)

    # left and right borders in red
    cv2.line(overlay, (x1, 0), (x1, h - 1), (0, 0, 255), 2)
    cv2.line(overlay, (x2, 0), (x2, h - 1), (0, 0, 255), 2)

    cv2.rectangle(overlay, (x1, y1), (x2, y2), (255, 255, 255), 2)

    cv2.putText(overlay, f"ROI: ({x1},{y1})-({x2},{y2})  {x2-x1}x{y2-y1}",
                (10, 80), cv2.FONT_HERSHEY_SIMPLEX, 0.6, (255, 255, 255), 2)

    cv2.imshow(WINDOW_NAME, overlay)

    roi = frame[y1:y2, x1:x2]
    if roi.size == 0 or roi.shape[0] < MIN_ROI_SIZE or roi.shape[1] < MIN_ROI_SIZE:
        key = cv2.waitKey(1) & 0xFF
        if key == ord('q'):
            break
        if key == ord('r'):
            reset_roi_to_center(w, h)
            prev_tensor = None
            print("ROI reset.")
        continue

    if roi_moved:
        prev_tensor = None
        roi_moved = False
        print("ROI moved -> baseline reset.")

    if time.time() - last_check >= CHECK_INTERVAL:
        last_check = time.time()

        curr_tensor = cv2_to_tensor(roi)

        # DEBUG: prove preproc is saving even without "change detected"
        if DEBUG_SAVE_PREPROC_EVERY_CHECK:
            ts_dbg = datetime.now().strftime("%Y-%m-%d_%H-%M-%S_%f")
            dbg_base = os.path.join(SAVE_FOLDER, f"dbg_{ts_dbg}")
            dbg_text, dbg_raw, dbg_pre = extract_text_and_save_preproc(roi, dbg_base)
            print(f"[DEBUG] Saved raw: {dbg_raw}")
            print(f"[DEBUG] Saved preproc: {dbg_pre}")
            print(f"[DEBUG] OCR: {dbg_text.strip()}")

        if prev_tensor is not None:
            dist = loss_fn(prev_tensor, curr_tensor).item()
            print(f"LPIPS Distance (ROI): {dist:.3f}")

            if dist > LPIPS_THRESHOLD:
                print(">>> CHANGE DETECTED (ROI) <<<")

                timestamp = datetime.now().strftime("%Y-%m-%d %H-%M-%S")
                base = os.path.join(SAVE_FOLDER, f"roi_change_{timestamp}")

                extracted_text, roi_raw_path, roi_preproc_path = extract_text_and_save_preproc(roi, base)

                print(f"Saved ROI raw: {roi_raw_path}")
                print(f"Saved ROI preprocessed: {roi_preproc_path}")
                print("Extracted text (ROI):")
                print(extracted_text)

                append_to_csv(CSV_FILE, timestamp, dist, extracted_text, roi_raw_path, roi_preproc_path)

                # Put the preprocessed image into Word (usually more useful than raw)
                write_to_word(timestamp, extracted_text, roi_preproc_path or roi_raw_path)
            else:
                print("No significant change (ROI)")

        prev_tensor = curr_tensor

    key = cv2.waitKey(1) & 0xFF
    if key == ord('q'):
        break
    if key == ord('r'):
        reset_roi_to_center(w, h)
        prev_tensor = None
        print("ROI reset.")

cap.release()
cv2.destroyAllWindows()